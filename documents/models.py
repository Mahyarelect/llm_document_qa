from django.db import models
from docx import Document as DocxDocument


class Document(models.Model):
    title = models.CharField(max_length=255)
    file = models.FileField(upload_to='documents/')
    full_text = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    def extract_text_from_docx(self):
        doc = DocxDocument(self.file.path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        return "\n".join(full_text)
    
    def create_chunks(self, chunk_size=1000, overlap=150):
        self.chunks.all().delete()

        text = self.full_text or ""
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end].strip()

            if chunk_text:
                DocumentChunk.objects.create(
                    document=self,
                    content=chunk_text,
                    chunk_index=chunk_index
                )

            chunk_index += 1
            start += chunk_size - overlap
            
    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)

        if self.file:
            extracted_text = self.extract_text_from_docx()

            if self.full_text != extracted_text:
                self.full_text = extracted_text
                super().save(update_fields=['full_text'])

            self.create_chunks()

    def __str__(self):
        return self.title


class DocumentChunk(models.Model):
    document = models.ForeignKey(
        Document,
        on_delete=models.CASCADE,
        related_name='chunks'
    )
    content = models.TextField()
    chunk_index = models.PositiveIntegerField()

    def __str__(self):
        return f"{self.document.title} - Chunk {self.chunk_index}"


class QuestionHistory(models.Model):
    question = models.TextField()
    answer = models.TextField(blank=True)
    retrieved_context = models.TextField(blank=True)
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.question[:80]